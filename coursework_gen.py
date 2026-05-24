import asyncio
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import google.generativeai as genai

try:
    from config import GEMINI_API_KEY
except ImportError:
    import os
    GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")

# Gemini modelini sozlash
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY.strip())
    gemini_model = genai.GenerativeModel("gemini-1.5-flash")


async def generate_coursework(topic: str, user_id: int) -> str:
    """Akademik va tizimli kurs ishini DOCX formatida yaratadi"""
    
    if not GEMINI_API_KEY:
        raise ValueError("Gemini API Key topilmadi. Iltimos, config faylini tekshiring.")

    prompt = (
        f"Siz nufuzli universitetning eng tajribali akademik professorisiz.\n"
        f"Mavzu: '{topic}' bo'yicha mukammal va to'liq universitet darajasidagi ilmiy kurs ishini o'zbek tilida yozing.\n\n"
        f"Struktura qat'iy ravishda quyidagi tartibda bo'lishi shart va har bir bo'limni juda keng, "
        f"akademik terminlar bilan chuqur yoritib bering (Xabarni qisqartirmang, imkon qadar ko'p ma'lumot yozing):\n\n"
        f"MUNDARIJA\n"
        f"KIRISH (Mavzuning dolzarbligi, maqsadi, vazifalari va ilmiy ahamiyati)\n"
        f"I-BO'LIM. [Bob sarlavhasi]\n"
        f"1.1. [Paragraph sarlavhasi va chuqur ilmiy nazariy tahlil]\n"
        f"1.2. [paragraph sarlavhasi va tahliliy tushuntirishlar]\n"
        f"II-BO'LIM. [Bob sarlavhasi]\n"
        f"2.1. [ paragraph sarlavhasi, amaliy tahlillar va statistik ma'lumotlar]\n"
        f"2.2. [ paragraph sarlavhasi va mavjud muammolarning yechimlari]\n"
        f"XULOSA (Kurs ishi davomida shakllantirilgan tavsiyalar va yakuniy xulosalar)\n"
        f"FOYDALANILGAN ADABIYOTLAR RO'YXATI (Kamida 5-8 ta ilmiy va darslik adabiyotlar, rasmiy manbalar va yili bilan)\n\n"
        f"DIQQAT: Markdown belgilari (masalan, *, **, # va hokazo) mutlaqo ishlatilmasin. Matn faqat toza akademik tilda bo'lsin."
    )

    try:
        response = await asyncio.to_thread(gemini_model.generate_content, prompt)
        full_text = response.text.strip()
    except Exception as e:
        raise RuntimeError(f"Gemini matn generatsiyasida xatolik: {e}")

    # DOCX faylini shakllantirish (Akademik OTM standarti)
    doc = Document()

    # Sahifa chetlarini belgilash (OTM standarti: Chap 3cm, O'ng 1.5cm, Tepa-past 2cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)  # ~2 cm
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.18)  # ~3 cm
        section.right_margin = Inches(0.59)  # ~1.5 cm

    # Times New Roman standard shrifti
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    # 1. Kurs ishi Tituli (Sarlavha sahifasi)
    p_univer = doc.add_paragraph()
    p_univer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_univer = p_univer.add_run("O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI\n\n\n")
    run_univer.bold = True
    run_univer.font.size = Pt(12)

    p_work = doc.add_paragraph()
    p_work.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_work = p_work.add_run("KURS ISHI\n\n")
    run_work.bold = True
    run_work.font.size = Pt(22)
    run_work.font.color.rgb = RGBColor(11, 29, 58)

    p_topic = doc.add_paragraph()
    p_topic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_topic = p_topic.add_run(f"MAVZU: \"{topic.upper()}\"\n\n\n\n\n")
    run_topic.bold = True
    run_topic.font.size = Pt(16)

    # Talaba haqida ma'lumot bloki (O'ng tomonga moslangan)
    p_info = doc.add_paragraph()
    p_info.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_info = p_info.add_run("Bajardi: ______________ (Talaba)\nIlmiy rahbar: ______________ (Professor)\n\n\n\n\n")
    run_info.font.size = Pt(12)
    run_info.italic = True

    # Shahar va Yil
    p_year = doc.add_paragraph()
    p_year.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_year = p_year.add_run("TOSHKENT - 2024")
    run_year.bold = True
    run_year.font.size = Pt(12)

    # Keyingi sahifaga o'tish (Page Break)
    doc.add_page_break()

    # 2. Kurs ishining asosiy qismini tahlil qilish va yozish
    lines = full_text.split("\n")
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue

        p = doc.add_paragraph()
        # Matnning ikki tomonini ham tekislash (Justify)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)

        # Sarlavhalarni aniqlash va ularni formatlash
        is_heading = False
        heading_words = ["MUNDARIJA", "KIRISH", "BO'LIM", "XULOSA", "ADABIYOTLAR", "1.", "2.", "3."]
        
        for hw in heading_words:
            if line_str.upper().startswith(hw):
                is_heading = True
                break
        
        if is_heading:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(line_str)
            run.bold = True
            run.font.size = Pt(15)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(8)
        else:
            # Oddiy paragraf (Abzats qo'shish)
            p.paragraph_format.first_line_indent = Inches(0.49)  # ~1.25 cm
            p.add_run(line_str)

    # Faylni xavfsiz va unikal nom bilan saqlash
    filename = f"coursework_{user_id}_{int(time.time())}.docx"
    doc.save(filename)
    return filename
